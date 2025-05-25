import pandas as pd
import os
from docx import Document

# ------------------------
# Load the Excel file
# ------------------------
file_path = "Band.xlsx"  # Adjust this path as needed
sheet_name = "Band"
df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)

# ------------------------
# Prepare output directory
# ------------------------
output_dir = "Grouped_Tables_By_Category_And_Question"
os.makedirs(output_dir, exist_ok=True)

# ------------------------
# Extract and group tables by (Category, Question)
# ------------------------
tables_by_pair = {}
i = 0

while i < len(df):
    if (
        pd.notna(df.iloc[i, 0]) and
        pd.notna(df.iloc[i + 1, 0]) and
        pd.isna(df.iloc[i + 2, 0])
    ):
        question_category = df.iloc[i, 0]
        question_text = df.iloc[i + 1, 0]
        group_key = (question_category, question_text)

        # Start of the table
        start_row = i + 3
        end_row = start_row

        # Find two consecutive empty rows
        while end_row + 1 < len(df):
            row_empty = df.iloc[end_row].isna().all()
            next_row_empty = df.iloc[end_row + 1].isna().all()
            if row_empty and next_row_empty:
                break
            end_row += 1

        # Extract table
        table = df.iloc[start_row:end_row + 1].reset_index(drop=True)

        # Add metadata
        metadata = pd.DataFrame([[question_category], [question_text], ['']])
        combined = pd.concat([metadata, table], ignore_index=True)

        # Group by (category, question)
        if group_key not in tables_by_pair:
            tables_by_pair[group_key] = []
        tables_by_pair[group_key].append(combined)

        # Move past the two empty rows
        i = end_row + 2
    else:
        i += 1

# ------------------------
# Save grouped tables to Excel files
# ------------------------
for idx, ((category, question), tables) in enumerate(tables_by_pair.items(), 1):
    safe_filename = f"Category_Question_{idx}.xlsx"
    output_path = os.path.join(output_dir, safe_filename)

    tables_with_gaps = []
    for table in tables:
        tables_with_gaps.append(table)
        tables_with_gaps.append(pd.DataFrame([[]]))  # Two empty rows

    combined_with_gaps = pd.concat(tables_with_gaps, ignore_index=True)
    combined_with_gaps.to_excel(output_path, index=False, header=False)

print(f"Saved grouped tables to '{output_dir}' directory.")

# ------------------------
# Save categorized questions to Word
# ------------------------
doc = Document()
doc.add_heading("List of Unique Survey Questions (with Categories)", level=1)

for idx, (category, question) in enumerate(tables_by_pair.keys(), 1):
    doc.add_paragraph(f"{idx}. [{category}]\n{question}")

doc_output_path = os.path.join(output_dir, "Unique_Categorized_Questions.docx")
doc.save(doc_output_path)

print(f"Saved unique categorized questions to: {doc_output_path}")
